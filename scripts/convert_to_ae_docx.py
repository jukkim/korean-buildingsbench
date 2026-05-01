"""
Convert paper_final.md to Applied Energy journal-format Word document.
Outputs: docs/paper_ae.docx

Applied Energy (Elsevier) requirements:
- Double-spaced, 12pt Times New Roman
- Line numbers (manual numbering note added)
- Title page / Highlights / Abstract / Keywords / Main text / References
- Figures embedded with captions (review format)
- Tables inline
"""

import re
import sys
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

BASE = Path(__file__).resolve().parent.parent
PAPER_MD = BASE / "docs" / "paper_final.md"
OUTPUT_DOCX = BASE / "docs" / "paper_ae.docx"
DOCS = BASE / "docs"

FIGURES = {
    "Fig. 1": DOCS / "fig1_pipeline.png",
    "Fig. 2": DOCS / "fig2_comparison.png",
    "Fig. 3": DOCS / "fig3_nscaling_new.png",
    "Fig. 4": DOCS / "fig4_revin_asymmetry.png",
}
FIGURE_CAPTIONS = {
    "Fig. 1": (
        "Fig. 1. End-to-end pipeline: from building archetype selection and 12D LHS parameter "
        "sampling through EnergyPlus simulation, Box-Cox normalization, RevIN-equipped Transformer "
        "training, to zero-shot inference without geographic information."
    ),
    "Fig. 2": (
        "Fig. 2. Main comparison of zero-shot commercial load forecasting performance (NRMSE, %) "
        "across six model configurations on the 955-building BuildingsBench evaluation set. "
        "The dashed line indicates the BB SOTA-M baseline (13.27%)."
    ),
    "Fig. 3": (
        "Fig. 3. N-scaling curve showing NRMSE as a function of the number of training buildings. "
        "Performance matches the SOTA from 70 buildings (n = 5) and saturates by 140 buildings "
        "(n = 10). The BB SOTA-M (13.27%) baseline is shown for reference."
    ),
    "Fig. 4": (
        "Fig. 4. RevIN's asymmetric effect across dataset scales. Green arrows indicate RevIN "
        "improvement (lower NRMSE); the red arrow indicates RevIN degradation on the full 900K "
        "BuildingsBench corpus."
    ),
}


# ── helpers ──────────────────────────────────────────────────────────────────

def set_double_spacing(para):
    pPr = para._p.get_or_add_pPr()
    lSpacing = OxmlElement("w:spacing")
    lSpacing.set(qn("w:line"), "480")      # 480 twips = double
    lSpacing.set(qn("w:lineRule"), "auto")
    lSpacing.set(qn("w:before"), "0")
    lSpacing.set(qn("w:after"), "0")
    pPr.append(lSpacing)


def set_single_spacing(para):
    pPr = para._p.get_or_add_pPr()
    lSpacing = OxmlElement("w:spacing")
    lSpacing.set(qn("w:line"), "240")
    lSpacing.set(qn("w:lineRule"), "auto")
    lSpacing.set(qn("w:before"), "0")
    lSpacing.set(qn("w:after"), "120")
    pPr.append(lSpacing)


def body_para(doc, text="", bold_spans=None, double=True):
    """Add a body paragraph with Times New Roman 12pt."""
    p = doc.add_paragraph()
    if double:
        set_double_spacing(p)
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)
    if bold_spans:
        # bold_spans: list of (start, end) positions in text
        pass
    return p


def add_formatted_para(doc, raw_text, double=True):
    """Parse inline bold (**text**) and italic (*text*) and add as a paragraph."""
    p = doc.add_paragraph()
    if double:
        set_double_spacing(p)

    parts = re.split(r"(\*\*.*?\*\*|\*[^*]+?\*|`[^`]+?`)", raw_text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = p.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("*") and part.endswith("*"):
            run = p.add_run(part[1:-1])
            run.italic = True
        elif part.startswith("`") and part.endswith("`"):
            run = p.add_run(part[1:-1])
            run.font.name = "Courier New"
        else:
            run = p.add_run(part)
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)
    return p


def add_heading(doc, text, level, double=True):
    """Add section heading."""
    p = doc.add_paragraph()
    if double:
        set_double_spacing(p)
    else:
        set_single_spacing(p)
    run = p.add_run(text)
    run.bold = True
    if level == 1:
        run.font.size = Pt(14)
    elif level == 2:
        run.font.size = Pt(13)
    else:
        run.font.size = Pt(12)
    run.font.name = "Times New Roman"
    pPr = p._p.get_or_add_pPr()
    spacing = pPr.find(qn("w:spacing"))
    if spacing is not None:
        spacing.set(qn("w:before"), "240")
    return p


def add_figure(doc, fig_key):
    """Embed figure image + caption."""
    img_path = FIGURES.get(fig_key)
    if img_path and img_path.exists():
        p_img = doc.add_paragraph()
        p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_single_spacing(p_img)
        run = p_img.add_run()
        run.add_picture(str(img_path), width=Inches(5.5))

        p_cap = doc.add_paragraph()
        set_single_spacing(p_cap)
        p_cap.alignment = WD_ALIGN_PARAGRAPH.LEFT
        caption_text = FIGURE_CAPTIONS.get(fig_key, fig_key)
        run_cap = p_cap.add_run(caption_text)
        run_cap.font.name = "Times New Roman"
        run_cap.font.size = Pt(10)
        run_cap.italic = True
        doc.add_paragraph()  # spacer


def add_markdown_table(doc, header_row, rows, caption=None):
    """Build a Word table from markdown table data."""
    ncols = len(header_row)
    nrows = len(rows) + 1  # +1 for header
    table = doc.add_table(rows=nrows, cols=ncols)
    table.style = "Table Grid"

    # header
    for j, cell_text in enumerate(header_row):
        cell = table.cell(0, j)
        cell.text = cell_text.strip()
        run = cell.paragraphs[0].runs[0]
        run.bold = True
        run.font.name = "Times New Roman"
        run.font.size = Pt(11)

    # data rows
    for i, row in enumerate(rows):
        for j, cell_text in enumerate(row):
            if j < ncols:
                cell = table.cell(i + 1, j)
                txt = cell_text.strip()
                # handle **bold** in cell
                parts = re.split(r"(\*\*.*?\*\*)", txt)
                para = cell.paragraphs[0]
                para.clear()
                for part in parts:
                    if part.startswith("**") and part.endswith("**"):
                        r = para.add_run(part[2:-2])
                        r.bold = True
                    else:
                        r = para.add_run(part)
                    r.font.name = "Times New Roman"
                    r.font.size = Pt(11)

    if caption:
        p = doc.add_paragraph()
        set_single_spacing(p)
        # Remove ** from caption for display
        clean_caption = re.sub(r"\*\*(.*?)\*\*", r"\1", caption)
        run = p.add_run(clean_caption)
        run.font.name = "Times New Roman"
        run.font.size = Pt(10)
        run.italic = True

    doc.add_paragraph()  # spacer after table


def page_break(doc):
    doc.add_page_break()


# ── main ──────────────────────────────────────────────────────────────────────

def parse_table_lines(lines):
    """Parse a list of markdown table lines into (header, rows)."""
    data_lines = [l for l in lines if not re.match(r"^\s*\|[-: |]+\|\s*$", l)]
    parsed = []
    for l in data_lines:
        cells = [c for c in l.strip().strip("|").split("|")]
        parsed.append(cells)
    if not parsed:
        return [], []
    return parsed[0], parsed[1:]


def build_document():
    doc = Document()

    # ── page layout: A4, 2.5cm margins ────────────────────────────────────
    section = doc.sections[0]
    section.page_height = Cm(29.7)
    section.page_width = Cm(21.0)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)

    md_lines = PAPER_MD.read_text(encoding="utf-8").splitlines()

    # ── TITLE PAGE ────────────────────────────────────────────────────────
    # Title
    title_text = (
        "Seven Hundred Simulations Suffice: Matching a 900,000-Building Foundation Model "
        "through Operational Diversity in Zero-Shot Load Forecasting"
    )
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_single_spacing(p_title)
    run_title = p_title.add_run(title_text)
    run_title.font.name = "Times New Roman"
    run_title.font.size = Pt(16)
    run_title.bold = True

    doc.add_paragraph()  # spacer

    # Author
    p_author = doc.add_paragraph()
    p_author.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_single_spacing(p_author)
    run_author = p_author.add_run("Jeong-Uk Kim")
    run_author.font.name = "Times New Roman"
    run_author.font.size = Pt(13)

    # Affiliation
    p_aff = doc.add_paragraph()
    p_aff.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_single_spacing(p_aff)
    run_aff = p_aff.add_run(
        "Department of Electrical Engineering, Sangmyung University, Seoul 03016, South Korea"
    )
    run_aff.font.name = "Times New Roman"
    run_aff.font.size = Pt(12)

    # Email
    p_email = doc.add_paragraph()
    p_email.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_single_spacing(p_email)
    run_email = p_email.add_run("E-mail: jukim@smu.ac.kr")
    run_email.font.name = "Times New Roman"
    run_email.font.size = Pt(12)

    doc.add_paragraph()  # spacer

    # Journal note
    p_journal = doc.add_paragraph()
    p_journal.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_single_spacing(p_journal)
    run_journal = p_journal.add_run("Submitted to: Applied Energy")
    run_journal.font.name = "Times New Roman"
    run_journal.font.size = Pt(12)
    run_journal.italic = True

    doc.add_paragraph()

    # Corresponding author note
    p_corr = doc.add_paragraph()
    p_corr.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_single_spacing(p_corr)
    run_corr = p_corr.add_run("Corresponding Author: Jeong-Uk Kim  jukim@smu.ac.kr")
    run_corr.font.name = "Times New Roman"
    run_corr.font.size = Pt(12)

    page_break(doc)

    # ── HIGHLIGHTS PAGE ───────────────────────────────────────────────────
    add_heading(doc, "Highlights", level=1, double=False)
    highlights = [
        "700 simulations match a 900,000-building foundation model in zero-shot forecasting",
        "12-dimensional Latin Hypercube Sampling generates operationally diverse training data",
        "RevIN helps small diverse data but degrades performance on large homogeneous corpora",
        "Geographic features are unnecessary: zero lat/lon yields identical accuracy",
        "Korean-trained model matches U.S. SOTA on real U.S. and Portuguese buildings",
    ]
    for h in highlights:
        p = doc.add_paragraph(style="List Bullet")
        set_single_spacing(p)
        run = p.add_run(h)
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)

    page_break(doc)

    # ── GRAPHICAL ABSTRACT ────────────────────────────────────────────────
    add_heading(doc, "Graphical Abstract", level=1, double=False)
    ga_path = DOCS / "graphical_abstract.png"
    if ga_path.exists():
        p_ga = doc.add_paragraph()
        p_ga.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_single_spacing(p_ga)
        p_ga.add_run().add_picture(str(ga_path), width=Inches(5.5))
    else:
        p_ga = doc.add_paragraph()
        run = p_ga.add_run("[Graphical Abstract — graphical_abstract.png]")
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)
        run.italic = True

    page_break(doc)

    # ── ABSTRACT ──────────────────────────────────────────────────────────
    abstract_text = (
        "Zero-shot building load forecasting---predicting energy consumption for unseen buildings "
        "without retraining---is essential for grid balancing, demand response, and energy management. "
        "BuildingsBench (NeurIPS 2023) assembled 900,000 U.S. building simulations and achieved 13.28% "
        "NRMSE on 955 commercial buildings (reproduced here at 13.27%), establishing scale as the prevailing "
        "paradigm. We challenge this with 700 EnergyPlus simulations---fifty per archetype across 14 building "
        "types---with operational schedules designed via 12-dimensional Latin Hypercube Sampling, combined with "
        "Reversible Instance Normalization (RevIN) and no geographic features. This small corpus achieves "
        "13.11 ± 0.17% NRMSE (best seed 12.93%) on the same evaluation set, matching and in the best-seed "
        "result surpassing the 900K-building SOTA. Both corpora train on single-country simulations yet "
        "evaluate on the same real buildings from four U.S. campuses and Portugal. Controlled experiments show "
        "the data-source advantage persists at 1.15 pp after controlling for augmentation; applying RevIN to "
        "the full 900K corpus degrades performance to 13.89%, consistent with our advantage stemming from data "
        "design rather than RevIN alone. The n-scaling analysis reveals a sharp transition: 70 buildings (five "
        "per archetype) already match the SOTA (13.28 ± 0.12%), with performance stabilizing from 140 "
        "buildings. Zero-shot evaluation on 218 real Korean convenience stores confirms sim-to-real transfer "
        "(12.30% vs. 13.14%). These results suggest several hundred operationally diverse parametric "
        "simulations can substitute for million-building corpora, substantially lowering the barrier to "
        "zero-shot forecasting in regions without large building stock databases."
    )

    add_heading(doc, "Abstract", level=1, double=False)
    p_abs = doc.add_paragraph()
    set_double_spacing(p_abs)
    run_abs = p_abs.add_run(abstract_text)
    run_abs.font.name = "Times New Roman"
    run_abs.font.size = Pt(12)

    doc.add_paragraph()

    # Keywords
    p_kw = doc.add_paragraph()
    set_single_spacing(p_kw)
    run_kw_label = p_kw.add_run("Keywords: ")
    run_kw_label.bold = True
    run_kw_label.font.name = "Times New Roman"
    run_kw_label.font.size = Pt(12)
    run_kw = p_kw.add_run(
        "building energy forecasting; zero-shot learning; foundation models; "
        "parametric simulation; data-centric AI; reversible instance normalization"
    )
    run_kw.font.name = "Times New Roman"
    run_kw.font.size = Pt(12)

    page_break(doc)

    # ── MAIN TEXT: parse MD ───────────────────────────────────────────────
    # We process from section 1 to Appendix C, skipping title/abstract already done
    sections_to_skip = {
        "abstract", "highlights", "keywords", "author", "affiliation", "e-mail"
    }

    i = 0
    in_table = False
    table_lines = []
    pending_table_caption = None

    # Sections we handle manually:
    skip_until_section = True  # skip until "## 1. Introduction"

    while i < len(md_lines):
        line = md_lines[i]

        # Detect start of main body
        if skip_until_section:
            if re.match(r"^## 1\.\s+Introduction", line):
                skip_until_section = False
            else:
                i += 1
                continue

        stripped = line.strip()

        # Page break markers (--- alone on a line)
        if stripped == "---":
            i += 1
            continue

        # Section heading H2
        m2 = re.match(r"^## (.+)$", line)
        if m2:
            heading_text = m2.group(1)
            add_heading(doc, heading_text, level=2)
            i += 1
            continue

        # Section heading H3
        m3 = re.match(r"^### (.+)$", line)
        if m3:
            heading_text = m3.group(1)
            add_heading(doc, heading_text, level=3)
            i += 1
            continue

        # Table start
        if stripped.startswith("|") and not in_table:
            in_table = True
            table_lines = [line]
            i += 1
            continue

        if in_table:
            if stripped.startswith("|"):
                table_lines.append(line)
                i += 1
                continue
            else:
                # Table ended — render it
                header, rows = parse_table_lines(table_lines)
                # Look ahead for caption (next **Table N.** line)
                caption = None
                if i < len(md_lines) and md_lines[i].strip().startswith("**Table"):
                    caption = md_lines[i].strip()
                    i += 1
                add_markdown_table(doc, header, rows, caption)
                in_table = False
                table_lines = []
                continue

        # Empty line
        if stripped == "":
            i += 1
            continue

        # Figure reference in "List of Figures" section → embed images
        m_fig = re.match(r"\*\*(Fig\. \d+)\.\*\*\s*(.*)", stripped)
        if m_fig:
            fig_key = m_fig.group(1)
            add_figure(doc, fig_key)
            i += 1
            continue

        # Math equation ($$...$$) — treat as centered italic
        if stripped.startswith("$$") and stripped.endswith("$$"):
            eq = stripped[2:-2].strip()
            p_eq = doc.add_paragraph()
            p_eq.alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_double_spacing(p_eq)
            run_eq = p_eq.add_run(eq)
            run_eq.font.name = "Times New Roman"
            run_eq.font.size = Pt(12)
            run_eq.italic = True
            i += 1
            continue

        # Multi-line math ($$\n...\n$$)
        if stripped == "$$":
            eq_lines = []
            i += 1
            while i < len(md_lines) and md_lines[i].strip() != "$$":
                eq_lines.append(md_lines[i].strip())
                i += 1
            i += 1  # skip closing $$
            eq_text = " ".join(eq_lines)
            p_eq = doc.add_paragraph()
            p_eq.alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_double_spacing(p_eq)
            run_eq = p_eq.add_run(eq_text)
            run_eq.font.name = "Times New Roman"
            run_eq.font.size = Pt(12)
            run_eq.italic = True
            continue

        # Skip "List of Figures" heading (we embedded figures already by content)
        if stripped == "## List of Figures":
            i += 1
            continue

        # Normal paragraph — with inline formatting
        add_formatted_para(doc, stripped, double=True)
        i += 1

    # ── FIGURE SECTION (after main text, before references) ───────────────
    # Figures are already embedded in the "List of Figures" section above,
    # but if any were in "List of Figures" we also add them here as standalone.

    # Actually, re-verify: figures in the List of Figures are already embedded.
    # The approach above handles it by matching **Fig. N.**

    page_break(doc)

    return doc


def main():
    print("Building Applied Energy format Word document...")
    doc = build_document()
    doc.save(str(OUTPUT_DOCX))
    print(f"Saved: {OUTPUT_DOCX}")


if __name__ == "__main__":
    main()
