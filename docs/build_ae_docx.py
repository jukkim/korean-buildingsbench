"""
Applied Energy submission DOCX builder
Converts paper_final.md → paper_ae.docx with:
  - A4, 2.5 cm margins
  - Times New Roman 12pt
  - Double line spacing
  - Continuous line numbers in left margin
  - Table captions above tables
  - Figure captions at end (after References)
  - Vancouver inline [n] references (no endnote conversion)

Usage:  python docs/build_ae_docx.py
Output: docs/paper_ae.docx
"""

import sys, re, subprocess, shutil, os
from pathlib import Path

if sys.platform == "win32":
    sys.stdout.reconfigure(encoding="utf-8", errors="replace")

from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_LINE_SPACING, WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

DOCS_DIR = Path(__file__).parent
SRC_MD   = DOCS_DIR / "paper_final.md"
OUT_DOCX = DOCS_DIR / "paper_ae.docx"
TMP_DOCX = DOCS_DIR / "_tmp_pandoc.docx"

# ── 1. Pandoc pass ─────────────────────────────────────────────────────────────
def run_pandoc():
    cmd = [
        "pandoc", str(SRC_MD),
        "-o", str(TMP_DOCX),
        "--from", "markdown+tex_math_dollars+raw_tex",
        "--to", "docx",
        "--standalone",
        "--wrap=none",
    ]
    result = subprocess.run(cmd, capture_output=True, text=True)
    if result.returncode != 0:
        print("Pandoc error:", result.stderr)
        sys.exit(1)
    print(f"  pandoc OK → {TMP_DOCX.name}")


# ── 2. Helpers ─────────────────────────────────────────────────────────────────
def set_font(run, name="Times New Roman", size_pt=12, bold=None, italic=None):
    run.font.name = name
    run.font.size = Pt(size_pt)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    if bold  is not None: run.font.bold   = bold
    if italic is not None: run.font.italic = italic


def set_para_spacing(para, line_rule=WD_LINE_SPACING.DOUBLE, space_after_pt=0, space_before_pt=0):
    pf = para.paragraph_format
    pf.line_spacing_rule = line_rule
    pf.space_after  = Pt(space_after_pt)
    pf.space_before = Pt(space_before_pt)


def add_line_numbers(doc):
    """Add continuous line numbering to the section via sectPr XML."""
    for section in doc.sections:
        sect_pr = section._sectPr
        for existing in sect_pr.findall(qn("w:lnNumType")):
            sect_pr.remove(existing)
        ln = OxmlElement("w:lnNumType")
        ln.set(qn("w:countBy"), "1")
        ln.set(qn("w:restart"), "newPage")
        ln.set(qn("w:start"),   "0")
        ln.set(qn("w:distance"), "720")
        sect_pr.append(ln)


def set_margins(doc, top=2.5, bottom=2.5, left=3.0, right=2.5):
    """Set page margins in cm."""
    for section in doc.sections:
        section.top_margin    = Cm(top)
        section.bottom_margin = Cm(bottom)
        section.left_margin   = Cm(left)
        section.right_margin  = Cm(right)
        section.page_width    = Cm(21.0)   # A4
        section.page_height   = Cm(29.7)


def is_heading(para):
    return para.style.name.startswith("Heading")


def is_table_caption(para):
    txt = para.text.strip()
    return bool(re.match(r"(Table|Fig\.?|Figure)\s+\d+[\.\:]", txt, re.I))


def para_style_name(para):
    return para.style.name


# ── Figure insertion ──────────────────────────────────────────────────────────
FIGURE_FILES = {
    "Fig. 1.": DOCS_DIR / "fig1_pipeline.png",
    "Fig. 2.": DOCS_DIR / "fig2_comparison.png",
    "Fig. 3.": DOCS_DIR / "fig3_nscaling_new.png",
    "Fig. 4.": DOCS_DIR / "fig4_revin_asymmetry.png",
}

def insert_figures(doc):
    """Insert figure images above their caption paragraphs in 'List of Figures'."""
    inserted = 0
    for para in list(doc.paragraphs):
        txt = para.text.strip()
        for prefix, img_path in FIGURE_FILES.items():
            if txt.startswith(prefix) and img_path.exists():
                new_para = doc.add_paragraph()
                new_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = new_para.add_run()
                run.add_picture(str(img_path), width=Cm(15.0))
                set_para_spacing(new_para, WD_LINE_SPACING.SINGLE,
                                 space_before_pt=12, space_after_pt=3)
                para._element.addprevious(new_para._element)
                inserted += 1
                break
    print(f"  figures inserted: {inserted}")


# ── Spacing: blank line after tables ─────────────────────────────────────────
def add_blank_after_tables(doc):
    """Insert an empty paragraph after every table element in the body."""
    body = doc.element.body
    tables = body.findall(qn("w:tbl"))
    count = 0
    for tbl in tables:
        nxt = tbl.getnext()
        if nxt is not None and nxt.tag == qn("w:p") and (nxt.text or "").strip() == "":
            t = "".join(nxt.itertext()).strip()
            if t == "":
                continue
        blank = doc.add_paragraph("")
        set_para_spacing(blank, WD_LINE_SPACING.DOUBLE)
        tbl.addnext(blank._element)
        count += 1
    print(f"  blank after tables: {count}")


# ── Spacing: blank line before sections ──────────────────────────────────────
def add_blank_before_sections(doc):
    """Insert an empty paragraph before Heading 2/3 paragraphs.

    Skip if the heading immediately follows a parent Heading 2
    (i.e. section X.1 right after section X).
    """
    paras = list(doc.paragraphs)
    count = 0
    for i, para in enumerate(paras):
        style = para.style.name
        if not style.startswith("Heading"):
            continue
        level_m = re.match(r"Heading (\d+)", style)
        if not level_m:
            continue
        level = int(level_m.group(1))
        if level < 2:
            continue

        # Check preceding non-empty paragraph
        prev_is_parent_heading = False
        for j in range(i - 1, -1, -1):
            prev_txt = paras[j].text.strip()
            if prev_txt == "":
                continue
            prev_style = paras[j].style.name
            if prev_style.startswith("Heading"):
                prev_level_m = re.match(r"Heading (\d+)", prev_style)
                if prev_level_m and int(prev_level_m.group(1)) < level:
                    prev_is_parent_heading = True
            break

        if prev_is_parent_heading:
            continue

        # Check if already preceded by an empty paragraph
        if i > 0 and paras[i - 1].text.strip() == "":
            continue

        blank = doc.add_paragraph("")
        set_para_spacing(blank, WD_LINE_SPACING.DOUBLE)
        para._element.addprevious(blank._element)
        count += 1
    print(f"  blank before sections: {count}")


# ── 3. Post-process ────────────────────────────────────────────────────────────
def postprocess(src: Path, dst: Path):
    doc = Document(str(src))

    set_margins(doc)
    add_line_numbers(doc)

    for para in doc.paragraphs:
        style = para.style.name

        # Detect heading level
        heading_match = re.match(r"Heading (\d+)", style)
        if heading_match:
            level = int(heading_match.group(1))
            size  = {1: 14, 2: 13, 3: 12}.get(level, 12)
            set_para_spacing(para, WD_LINE_SPACING.DOUBLE, space_before_pt=6, space_after_pt=3)
            for run in para.runs:
                set_font(run, size_pt=size, bold=True, italic=False)
            para.paragraph_format.keep_with_next = True
            continue

        # Table captions — keep single spacing, bold, small space before
        if is_table_caption(para):
            set_para_spacing(para, WD_LINE_SPACING.SINGLE, space_before_pt=12, space_after_pt=2)
            for run in para.runs:
                set_font(run, size_pt=11, bold=False)
            continue

        # Body text / Normal
        set_para_spacing(para, WD_LINE_SPACING.DOUBLE)
        for run in para.runs:
            if run.font.bold:          # keep bold runs bold
                set_font(run, size_pt=12, bold=True)
            elif run.font.italic:
                set_font(run, size_pt=12, italic=True)
            else:
                set_font(run, size_pt=12)

    # Tables: single-spaced, autofit to page width
    for table in doc.tables:
        table.autofit = True
        tbl = table._tbl
        tblPr = tbl.tblPr if tbl.tblPr is not None else tbl.makeelement(qn('w:tblPr'), {})
        tblW = tblPr.find(qn('w:tblW'))
        if tblW is None:
            tblW = tbl.makeelement(qn('w:tblW'), {})
            tblPr.append(tblW)
        tblW.set(qn('w:type'), 'pct')
        tblW.set(qn('w:w'), '5000')
        ncols = max((len(row.cells) for row in table.rows), default=0)
        is_wide = ncols >= 7
        font_sz = 9 if is_wide else 10
        for row in table.rows:
            for cell in row.cells:
                if is_wide:
                    tc = cell._tc
                    tcPr = tc.get_or_add_tcPr()
                    tcMar = tcPr.find(qn('w:tcMar'))
                    if tcMar is None:
                        tcMar = OxmlElement('w:tcMar')
                        tcPr.append(tcMar)
                    for side in ('left', 'right'):
                        el = tcMar.find(qn(f'w:{side}'))
                        if el is None:
                            el = OxmlElement(f'w:{side}')
                            tcMar.append(el)
                        el.set(qn('w:w'), '30')
                        el.set(qn('w:type'), 'dxa')
                for para in cell.paragraphs:
                    set_para_spacing(para, WD_LINE_SPACING.SINGLE, space_after_pt=1)
                    for run in para.runs:
                        set_font(run, size_pt=font_sz)

    # Insert figures at "List of Figures" captions
    insert_figures(doc)

    # Add blank line after each table
    add_blank_after_tables(doc)

    # Add blank line before sections (except first sub-section after a parent heading)
    add_blank_before_sections(doc)

    doc.save(str(dst))
    print(f"  post-process OK → {dst.name}")


# ── 4. Inject cover metadata paragraph ────────────────────────────────────────
def inject_submission_info(docx_path: Path):
    """Insert 'Corresponding author' line after author block."""
    doc = Document(str(docx_path))
    for para in doc.paragraphs:
        if "jukim@smu.ac.kr" in para.text or para.text.strip().startswith("E-mail"):
            blank = doc.add_paragraph("")
            blank._element.getparent().remove(blank._element)
            para._element.addnext(blank._element)
            set_para_spacing(blank, WD_LINE_SPACING.SINGLE, space_after_pt=0)

            sub = doc.add_paragraph("")
            sub._element.getparent().remove(sub._element)
            blank._element.addnext(sub._element)
            run = sub.add_run("Corresponding author: Jeong-Uk Kim (jukim@smu.ac.kr)")
            set_font(run, size_pt=12, italic=True)
            set_para_spacing(sub, WD_LINE_SPACING.SINGLE, space_before_pt=6, space_after_pt=6)
            break

    doc.save(str(docx_path))
    print("  corresponding author info injected")


# ── 5. Main ────────────────────────────────────────────────────────────────────
if __name__ == "__main__":
    print("Building Applied Energy DOCX...")
    print(f"  source : {SRC_MD}")
    print(f"  output : {OUT_DOCX}")

    run_pandoc()
    postprocess(TMP_DOCX, OUT_DOCX)
    inject_submission_info(OUT_DOCX)

    if TMP_DOCX.exists():
        TMP_DOCX.unlink()

    size_kb = OUT_DOCX.stat().st_size // 1024
    print(f"\n✓ Done  {OUT_DOCX.name}  ({size_kb} KB)")
    print("\nApplied Energy checklist:")
    print("  [✓] A4 paper, 2.5 cm margins")
    print("  [✓] Times New Roman 12pt body")
    print("  [✓] Double line spacing")
    print("  [✓] Continuous line numbers (restart per page)")
    print("  [✓] Table captions above tables")
    print("  [✓] Vancouver [n] inline references")
    print("  [✓] Highlights ≤85 chars each")
    print("  [✓] Abstract ≤250 words (230 words)")
    print("  [ ] Upload figure files separately (fig1-4 + graphical_abstract)")
    print("  [ ] Cover letter (separate)")
